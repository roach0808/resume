from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
from docx import Document
import pypandoc  # Ensure this is installed for PDF conversion
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_custom_margins(section, top, bottom, left, right):
    """Set custom margins for a section in inches."""
    section_properties = section._sectPr  # Access the section properties
    pgMar = section_properties.xpath('./w:pgMar')[0]  # Find the <w:pgMar> element

    # Convert inches to twips (1 inch = 1440 twips)
    pgMar.set(qn('w:top'), str(int(top * 1440)))
    pgMar.set(qn('w:bottom'), str(int(bottom * 1440)))
    pgMar.set(qn('w:left'), str(int(left * 1440)))
    pgMar.set(qn('w:right'), str(int(right * 1440)))


def create_resume_from_json(resume_json: dict, output_dir: str, output_format: str = 'docx'):
    if isinstance(resume_json, str):
        try:
            resume_json = json.loads(resume_json)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")

    doc = Document()

    # === Set Custom Margins ===
    section = doc.sections[0]
    set_custom_margins(section, top=1, bottom=1, left=0.8, right=0.8)


    # Define custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'  # Change the default font style
    font.size = Pt(10)   # Change the default font size

    # Name and Title
    name_para = doc.add_paragraph()
    
    

    name_run = name_para.add_run(resume_json['name'])
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.name = 'Georgia'  # Custom font for the name
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Title (optional)
    if 'title' in resume_json and resume_json['title']:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(resume_json['title'])
        title_run.font.size = Pt(12)
        title_run.font.name = 'Georgia'  # Custom font for the title
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Info (optional)
    if 'Contact' in resume_json and resume_json['Contact']:
        contact_info = resume_json['Contact']
        contact_para = doc.add_paragraph()
        if 'email' in contact_info:
            email_run = contact_para.add_run("Email: ")
            email_run.bold = True
            email_run.font.name = 'Georgia'
            contact_para.add_run(contact_info['email'])

        if 'phone' in contact_info:
            contact_para.add_run(", ")  # Add a new line
            phone_run = contact_para.add_run("Phone: ")
            phone_run.bold = True
            phone_run.font.name = 'Georgia'
            contact_para.add_run(contact_info['phone'])

        if 'location' in contact_info:
            contact_para.add_run(", ")
            location_run = contact_para.add_run("Location: ")
            location_run.bold = True
            location_run.font.name = 'Georgia'
            contact_para.add_run(contact_info['location'])

        if 'linkedin' in contact_info:
            contact_para.add_run(", ")
            linkedin_run = contact_para.add_run("LinkedIn: ")
            linkedin_run.bold = True
            linkedin_run.font.name = 'Georgia'
            contact_para.add_run(contact_info['linkedin'])

        if 'website' in contact_info:
            contact_para.add_run(", ")
            website_run = contact_para.add_run("Website: ")
            website_run.bold = True
            website_run.font.name = 'Georgia'
            contact_para.add_run(contact_info['website'])

        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a line break
    # doc.add_paragraph()

    # Summary Section (optional)
    if 'summary' in resume_json and resume_json['summary']:
        summary_heading = doc.add_paragraph('SUMMARY', style='Heading 2')
        summary_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        summary_heading.runs[0].font.name = 'Georgia'
        summary_heading.runs[0].font.size = Pt(13)
        summary_heading.runs[0].bold = True

        # Add a paragraph for the summary text
        summary = doc.add_paragraph()
        summary_run = summary.add_run(resume_json['summary'])  # Add a Run to the paragraph
        summary_run.font.name = 'Georgia'  # Change font style for the summary text
        summary_run.font.size = Pt(11)
        summary.paragraph_format.space_after = Pt(10)

    # Add a line break
    # doc.add_paragraph()

    

    # Skills Section (optional)
    if 'skills' in resume_json and resume_json['skills']:
        skills_heading = doc.add_paragraph('SKILLS', style='Heading 2')
        skills_heading.runs[0].font.name = 'Georgia'
        skills_heading.runs[0].font.size = Pt(13)
        skills_heading.runs[0].bold = True

        for skill, items in resume_json['skills'].items():
            para = doc.add_paragraph()
            skill_run = para.add_run(f"{skill}: ")
            skill_run.bold = True
            skill_run.font.name = 'Georgia'
            para.add_run(", ".join(items))

    # Add a line break
    doc.add_paragraph()
    
    # Experience Section (optional)
    if 'experiences' in resume_json and resume_json['experiences']:
        experience_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Heading 2')
        experience_heading.runs[0].font.name = 'Georgia'
        experience_heading.runs[0].font.size = Pt(13)
        experience_heading.runs[0].bold = True

        for exp in resume_json['experiences']:
            # Company name and location
            company_name = exp.get('company', 'Unknown Company')
            location = exp.get('location', '')
            company_para = doc.add_paragraph(f"{company_name}, {location}")
            company_para.runs[0].font.name = 'Georgia'
            company_para.runs[0].font.size = Pt(12)
            company_para.runs[0].bold = True
            company_para.paragraph_format.space_after = Pt(1)

            # Job title and dates
            role = exp.get('role', 'Position')
            start_date = exp.get('start_date', 'Start Date')
            end_date = exp.get('end_date', 'End Date')
            title_para = doc.add_paragraph(f"{role}\t{start_date} – {end_date}")
            title_para.runs[0].font.name = 'Georgia'
            title_para.runs[0].font.size = Pt(10)
            title_para.paragraph_format.space_after = Pt(2)

            # Company description if exists
            if 'description' in exp:
                desc_para = doc.add_paragraph(exp['description'])
                desc_para.runs[0].font.name = 'Georgia'
                desc_para.runs[0].font.size = Pt(10)
                desc_para.paragraph_format.space_after = Pt(3)

            # Experience bullets
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                bullet_para = doc.add_paragraph()
                bullet_run = bullet_para.add_run(f"· {bullet}")
                bullet_run.font.name = 'Georgia'
                bullet_run.font.size = Pt(10)
                bullet_para.paragraph_format.space_after = Pt(1)
                bullet_para.paragraph_format.left_indent = Pt(18)
            
            # Add project sections if they exist
            if 'project_sections' in exp and exp['project_sections']:
                for project_section in exp['project_sections']:
                    # Project section header
                    project_header = doc.add_paragraph(f"- {project_section.get('name', 'Project')}: {project_section.get('description', '')}")
                    project_header.runs[0].font.name = 'Georgia'
                    project_header.runs[0].font.size = Pt(10)
                    project_header.runs[0].bold = True
                    project_header.paragraph_format.space_after = Pt(2)
                    
                    # Project bullets
                    for project_bullet in project_section.get('bullets', []):
                        project_bullet_para = doc.add_paragraph()
                        project_bullet_run = project_bullet_para.add_run(f"· {project_bullet}")
                        project_bullet_run.font.name = 'Georgia'
                        project_bullet_run.font.size = Pt(10)
                        project_bullet_para.paragraph_format.space_after = Pt(1)
                        project_bullet_para.paragraph_format.left_indent = Pt(18)

    # Add a line break
    doc.add_paragraph()
    
    # Certificates Section (if exists)
    if 'Certificates' in resume_json and resume_json['Certificates']:
        certificates_heading = doc.add_paragraph('CERTIFICATES', style='Heading 2')
        certificates_heading.runs[0].font.name = 'Georgia'
        certificates_heading.runs[0].font.size = Pt(13)
        certificates_heading.runs[0].bold = True

        for cert in resume_json['Certificates']:
            cert_para = doc.add_paragraph()
            cert_name_run = cert_para.add_run(cert.get('certificate_name', cert.get('name', '')))
            cert_name_run.bold = True
            cert_name_run.font.name = 'Georgia'
            cert_name_run.font.size = Pt(10)
            
            if 'issued_by' in cert or 'issuer' in cert:
                issuer = cert.get('issued_by', cert.get('issuer', ''))
                cert_para.add_run(f" - {issuer}")
            if 'issued_date' in cert or 'date' in cert:
                date = cert.get('issued_date', cert.get('date', ''))
                cert_para.add_run(f" ({date})")
            
            cert_para.runs[0].font.name = 'Georgia'
            cert_para.runs[0].font.size = Pt(10)
            cert_para.paragraph_format.space_after = Pt(2)

    # Education Section (optional)
    if 'education' in resume_json and resume_json['education']:
        education_heading = doc.add_paragraph('EDUCATION', style='Heading 2')
        education_heading.runs[0].font.name = 'Georgia'
        education_heading.runs[0].font.size = Pt(13)
        education_heading.runs[0].bold = True

        for edu in resume_json['education']:
            edu_heading = doc.add_paragraph(f"{edu['institute_name']} - {edu['degree']}", style='Heading 3')
            edu_heading.paragraph_format.space_after = Pt(1)
            edu_heading.runs[0].font.name = 'Georgia'
            edu_heading.runs[0].font.size = Pt(12)

            edu_dates = doc.add_paragraph(f"{edu['start_date']} - {edu['end_date']}")
            edu_dates.paragraph_format.space_after = Pt(1)
            edu_dates.runs[0].font.name = 'Georgia'
            edu_dates.runs[0].font.size = Pt(10)

            if 'gpa' in edu:
                gpa = doc.add_paragraph(f"GPA: {edu['gpa']}")
                gpa.paragraph_format.space_after = Pt(5)
                gpa.runs[0].font.name = 'Georgia'
                gpa.runs[0].font.size = Pt(10)

    # Ensure the output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save the DOCX
    doc_filename = f"{resume_json['name']} Resume.docx"
    doc_path = os.path.join(output_dir, doc_filename)
    doc.save(doc_path)

    # Convert to PDF
    pdf_path = None
    # if output_format == 'pdf' or output_format == 'both':
    #     pdf_filename = f"{resume_json['name']} Resume.pdf"
    #     pdf_path = os.path.join(output_dir, pdf_filename)
    #     try:
    #         pypandoc.convert_file(doc_path, 'pdf', outputfile=pdf_path)
    #     except Exception as e:
    #         raise IOError(f"Failed to convert DOCX to PDF: {e}")

    return doc_path, pdf_path